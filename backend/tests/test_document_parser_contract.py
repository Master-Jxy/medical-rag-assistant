"""Stage 24.1: normalized document parser contract and compatibility."""

from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document as DocxDocument

from app.core.exceptions import DocumentParseError
from app.modules.knowledge.ingestion import (
    FileTypePolicy,
    ParseQuality,
    ParseRequest,
    ParsedDocument,
    ParsedElement,
    ParserRegistration,
    ParserRegistry,
)
from app.modules.knowledge.parser import LocalDocumentParser, ParsedPreview


REQUIRED_DOCX_ENTRIES = {
    "[Content_Types].xml": b"<Types />",
    "word/document.xml": b"<w:document />",
}


def write_docx(path: Path) -> bytes:
    document = DocxDocument()
    document.add_heading("护理计划", level=1)
    document.add_paragraph("每日记录症状变化。")
    document.add_paragraph("按时复诊", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "血压"
    table.cell(1, 1).text = "每日测量"
    buffer = BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    path.write_bytes(data)
    return data


class DummyParser:
    def __init__(self, text: str) -> None:
        self.text = text
        self.called = False

    def parse(self, request: ParseRequest) -> ParsedDocument:
        self.called = True
        return ParsedDocument(
            document_metadata={
                "file_name": request.file_name or request.path.name,
                "page_count": 1,
            },
            elements=(
                ParsedElement(
                    element_id="dummy-1",
                    kind="paragraph",
                    text=self.text,
                    page_no=1,
                    order=1,
                ),
            ),
            quality=ParseQuality(status="pass", counts={"text": 1}),
        )


def test_registry_selects_by_explicit_suffix_and_rejects_unknown(tmp_path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("医学资料", encoding="utf-8")
    txt_parser = DummyParser("TXT")
    pdf_parser = DummyParser("PDF")
    registry = ParserRegistry(
        [
            ParserRegistration("pdf", pdf_parser, suffixes=(".pdf",)),
            ParserRegistration("txt", txt_parser, suffixes=(".txt",)),
        ]
    )

    parsed = registry.parse(ParseRequest(path=path, suffix=".TXT"))

    assert parsed.text == "TXT"
    assert txt_parser.called is True
    assert pdf_parser.called is False
    with pytest.raises(DocumentParseError) as exc_info:
        registry.parse(ParseRequest(path=Path("missing.docx"), suffix=".docx"))
    assert exc_info.value.code == "DOCUMENT_PARSE_ERROR"


def test_local_parser_returns_normalized_document_and_legacy_preview(tmp_path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("第一段\n\n第二段", encoding="utf-8")
    parser = LocalDocumentParser()

    document = parser.parse_document(ParseRequest(path=path, suffix=".txt"))
    preview = parser.parse(path, ".txt")

    assert document.document_metadata["parser"] == "local_pdf_txt"
    assert [element.text for element in document.elements] == ["第一段\n\n第二段"]
    assert document.assets == ()
    assert document.quality.status == "pass"
    assert document.quality.counts == {
        "text": 1,
        "table_like": 0,
        "scanned_or_image": 0,
    }
    assert preview == ParsedPreview(
        "第一段\n\n第二段",
        1,
        (),
        {
            "status": "pass",
            "counts": {
                "text": 1,
                "table_like": 0,
                "scanned_or_image": 0,
            },
            "page_results": [
                {"page": 1, "kind": "text", "text_chars": 8, "image_count": 0}
            ],
        },
    )


def test_file_type_policy_rejects_spoofed_binary_and_damaged_docx(tmp_path) -> None:
    fake_pdf = tmp_path / "fake.pdf"
    fake_pdf.write_text("not a pdf", encoding="utf-8")
    with pytest.raises(DocumentParseError, match="PDF"):
        FileTypePolicy.validate_path(fake_pdf, ".pdf")

    text_with_nul = tmp_path / "bad.md"
    text_with_nul.write_bytes(b"title\x00body")
    with pytest.raises(DocumentParseError, match="空字节"):
        FileTypePolicy.validate_path(text_with_nul, ".md")

    damaged_docx = tmp_path / "bad.docx"
    damaged_docx.write_bytes(b"not a zip")
    with pytest.raises(DocumentParseError, match="DOCX"):
        FileTypePolicy.validate_path(damaged_docx, ".docx")


def write_minimal_docx(path: Path, entries: dict[str, bytes]) -> None:
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as package:
        for name, content in entries.items():
            package.writestr(name, content)


def test_docx_policy_rejects_high_compression_ratio(tmp_path) -> None:
    path = tmp_path / "ratio.docx"
    entries = {
        **REQUIRED_DOCX_ENTRIES,
        "word/large.xml": b"A" * (1024 * 1024),
    }
    write_minimal_docx(path, entries)

    with pytest.raises(DocumentParseError, match="压缩比"):
        FileTypePolicy.validate_path(path, ".docx")


def test_docx_policy_rejects_metadata_before_reading_entries(tmp_path, monkeypatch) -> None:
    path = tmp_path / "metadata.docx"
    path.write_bytes(b"placeholder")

    def run_with_infos(infos, expected: str) -> None:
        class FakeZip:
            def __init__(self, _: Path) -> None:
                pass

            def __enter__(self):
                return self

            def __exit__(self, *args) -> None:
                return None

            def infolist(self):
                return infos

            def read(self, _info):
                raise AssertionError("entry content should not be read before metadata passes")

        monkeypatch.setattr(
            "app.modules.knowledge.ingestion.file_types.is_zipfile",
            lambda _: True,
        )
        monkeypatch.setattr("app.modules.knowledge.ingestion.file_types.ZipFile", FakeZip)
        with pytest.raises(DocumentParseError, match=expected):
            FileTypePolicy.validate_path(path, ".docx")

    normal = [
        SimpleNamespace(
            filename=name,
            file_size=len(content),
            compress_size=max(1, len(content)),
            flag_bits=0,
        )
        for name, content in REQUIRED_DOCX_ENTRIES.items()
    ]
    run_with_infos(
        [
            SimpleNamespace(
                filename=f"word/item-{index}.xml",
                file_size=1,
                compress_size=1,
                flag_bits=0,
            )
            for index in range(257)
        ],
        "过多",
    )
    run_with_infos(
        normal
        + [
            SimpleNamespace(
                filename="word/big.xml",
                file_size=51 * 1024 * 1024,
                compress_size=51 * 1024 * 1024,
                flag_bits=0,
            )
        ],
        "单个文件条目过大",
    )
    run_with_infos(
        normal
        + [
            SimpleNamespace(
                filename=f"word/part-{index}.xml",
                file_size=10 * 1024 * 1024,
                compress_size=10 * 1024 * 1024,
                flag_bits=0,
            )
            for index in range(6)
        ],
        "总体积过大",
    )
    run_with_infos(
        normal
        + [
            SimpleNamespace(
                filename="word/encrypted.xml",
                file_size=1,
                compress_size=1,
                flag_bits=0x1,
            )
        ],
        "加密",
    )
    run_with_infos(
        normal
        + [
            SimpleNamespace(
                filename="../evil.xml",
                file_size=1,
                compress_size=1,
                flag_bits=0,
            )
        ],
        "路径异常",
    )
    run_with_infos(
        normal
        + [
            SimpleNamespace(
                filename="_rels/.rels",
                file_size=300 * 1024,
                compress_size=300 * 1024,
                flag_bits=0,
            )
        ],
        "关系文件过大",
    )


def test_docx_policy_rejects_external_relationship_with_single_quotes(tmp_path) -> None:
    path = tmp_path / "external.docx"
    write_minimal_docx(
        path,
        {
            **REQUIRED_DOCX_ENTRIES,
            "_rels/.rels": b"""
            <Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>
              <Relationship Id='rId1' Type='x' Target='https://example.com' TargetMode='External'/>
            </Relationships>
            """,
        },
    )

    with pytest.raises(DocumentParseError, match="外部关系"):
        FileTypePolicy.validate_path(path, ".docx")


def test_docx_policy_rejects_malformed_relationship_xml(tmp_path) -> None:
    path = tmp_path / "bad-rels.docx"
    write_minimal_docx(
        path,
        {
            **REQUIRED_DOCX_ENTRIES,
            "_rels/.rels": b"<Relationships><Relationship TargetMode='External'",
        },
    )

    with pytest.raises(DocumentParseError, match="关系XML无效"):
        FileTypePolicy.validate_path(path, ".docx")


def test_docx_parser_preserves_title_list_and_table_semantics(tmp_path) -> None:
    path = tmp_path / "sample.docx"
    write_docx(path)
    FileTypePolicy.validate_path(path, ".docx")

    parsed = LocalDocumentParser().parse_document(ParseRequest(path=path, suffix=".docx"))

    assert [element.kind for element in parsed.elements] == [
        "title",
        "paragraph",
        "list",
        "table",
    ]
    assert "护理计划" in parsed.text
    assert parsed.elements[-1].table_html.startswith("<table>")
    assert parsed.quality.counts["table"] == 1


def test_markdown_parser_preserves_heading_list_and_table_semantics(tmp_path) -> None:
    path = tmp_path / "sample.md"
    path.write_text(
        "# 随访记录\n\n- 记录血压\n\n| 指标 | 结果 |\n| --- | --- |\n| 血压 | 稳定 |\n",
        encoding="utf-8",
    )
    FileTypePolicy.validate_path(path, ".md")

    parsed = LocalDocumentParser().parse_document(ParseRequest(path=path, suffix=".md"))

    assert [element.kind for element in parsed.elements] == ["title", "list", "table"]
    assert parsed.elements[0].text == "随访记录"
    assert "血压 | 稳定" in parsed.elements[-1].text


def test_html_parser_removes_dangerous_nodes_and_keeps_body_semantics(tmp_path) -> None:
    path = tmp_path / "sample.html"
    path.write_text(
        """
        <html><head><style>body{display:none}</style><script>alert(1)</script></head>
        <body><h1>康复资料</h1><p>安全正文</p><form><input value="secret"></form>
        <iframe src="https://example.com"></iframe><table onclick="evil()"><tr><td>项目</td><td>说明</td></tr></table></body></html>
        """,
        encoding="utf-8",
    )
    FileTypePolicy.validate_path(path, ".html")

    parsed = LocalDocumentParser().parse_document(ParseRequest(path=path, suffix=".html"))

    assert [element.kind for element in parsed.elements] == ["title", "paragraph", "table"]
    assert "alert" not in parsed.text
    assert "secret" not in parsed.text
    assert "iframe" not in parsed.text
    assert "onclick" not in (parsed.elements[-1].table_html or "")


def test_legacy_preview_is_projected_from_structured_document() -> None:
    document = ParsedDocument(
        document_metadata={"page_count": 1},
        elements=(
            ParsedElement(
                element_id="e1",
                kind="paragraph",
                text="x" * 8001,
                page_no=1,
                order=1,
            ),
        ),
        warnings=("source warning",),
        quality=ParseQuality(status="warning", counts={"text": 1}),
    )

    preview = ParsedPreview.from_document(document)

    assert len(preview.text) == 8000
    assert preview.warnings == ("source warning", "预览已截断")
    assert preview.quality == {
        "status": "warning",
        "counts": {"text": 1},
        "page_results": [],
    }
