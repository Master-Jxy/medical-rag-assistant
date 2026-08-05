"""无模型文档预解析Port与本地适配器。"""

import hashlib
from dataclasses import dataclass
from html import escape
from pathlib import Path
from threading import BoundedSemaphore
from typing import Protocol

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from markdown_it import MarkdownIt
from PIL import Image, UnidentifiedImageError
from pypdf import PdfReader

from app.core.exceptions import DocumentParseError
from app.modules.knowledge.ingestion import (
    DocumentParserPort,
    FileTypePolicy,
    ParseQuality,
    ParseRequest,
    ParsedAsset,
    ParsedDocument,
    ParsedElement,
    ParserRegistration,
    ParserRegistry,
)


@dataclass(frozen=True, slots=True)
class ParsedPreview:
    text: str
    page_count: int
    warnings: tuple[str, ...] = ()
    quality: dict[str, object] | None = None

    @classmethod
    def from_document(cls, document: ParsedDocument) -> "ParsedPreview":
        text = document.text
        warnings = list(document.warnings)
        quality = document.quality.to_legacy_dict()
        if len(text) > 8000:
            warnings.append("预览已截断")
        if document.assets:
            quality["assets"] = [
                {
                    "asset_id": asset.asset_id,
                    "kind": asset.kind,
                    "page_no": asset.page_no,
                    "mime_type": asset.mime_type,
                    "sha256": asset.sha256,
                    "metadata": {
                        key: value
                        for key, value in dict(asset.metadata).items()
                        if key not in {"storage_ref", "path"}
                    },
                }
                for asset in document.assets
            ]
        enrichment = document.document_metadata.get("enrichment")
        if isinstance(enrichment, dict):
            quality["enrichment"] = dict(enrichment)
        return cls(
            text=text[:8000],
            page_count=document.page_count,
            warnings=tuple(dict.fromkeys(warnings)),
            quality=quality,
        )


class ParserPort(Protocol):
    def parse(self, path: Path, suffix: str) -> ParsedPreview: ...


class PdfCandidateFallbackParser:
    name = "pdf_candidate_fallback"
    _candidate_semaphore = BoundedSemaphore(1)

    def __init__(
        self,
        baseline: DocumentParserPort,
        candidate: DocumentParserPort | None,
        *,
        enabled: bool = False,
        promoted: bool = False,
        max_pages: int = 20,
        max_file_size_bytes: int = 10 * 1024 * 1024,
    ) -> None:
        self.baseline = baseline
        self.candidate = candidate
        self.enabled = enabled
        self.promoted = promoted
        self.max_pages = max_pages
        self.max_file_size_bytes = max_file_size_bytes

    def parse(self, request: ParseRequest) -> ParsedDocument:
        if request.normalized_suffix != ".pdf" or not self.enabled:
            return self.baseline.parse(request)
        baseline = self.baseline.parse(request)
        warning_prefix = "Docling候选已回退PyPDF："
        try:
            self._check_resource_bounds(request.path)
            if self.candidate is None:
                raise DocumentParseError("候选解析器不可用")
            if not self._candidate_semaphore.acquire(blocking=False):
                raise DocumentParseError("候选解析并发限制已满")
            try:
                candidate = self.candidate.parse(request)
            finally:
                self._candidate_semaphore.release()
            self._validate_candidate(candidate, baseline)
            if not self.promoted:
                return _with_warning(baseline, "Docling候选仅完成实验观测，尚未批准替换PyPDF")
            return _with_warning(candidate, "Docling候选已批准替换PyPDF")
        except Exception as exc:
            reason = str(exc).strip() or type(exc).__name__
            return _with_warning(baseline, f"{warning_prefix}{reason}")

    def _check_resource_bounds(self, path: Path) -> None:
        if path.stat().st_size > self.max_file_size_bytes:
            raise DocumentParseError("文件超过候选解析大小限制")
        try:
            page_count = len(PdfReader(str(path)).pages)
        except Exception as exc:
            raise DocumentParseError("无法读取PDF页数") from exc
        if page_count > self.max_pages:
            raise DocumentParseError("页数超过候选解析限制")

    @staticmethod
    def _validate_candidate(candidate: ParsedDocument, baseline: ParsedDocument) -> None:
        if not candidate.elements or not candidate.text.strip():
            raise DocumentParseError("候选解析输出为空")
        pages = [
            element.page_no
            for element in candidate.elements
            if element.page_no is not None
        ]
        if pages and pages != sorted(pages):
            raise DocumentParseError("候选解析页序异常")
        if candidate.page_count and baseline.page_count and candidate.page_count != baseline.page_count:
            raise DocumentParseError("候选解析页数与基线不一致")
        status = candidate.quality.status
        if status == "failed":
            raise DocumentParseError("候选解析质量未通过")


class LocalStructuredDocumentParser:
    name = "local_pdf_txt"

    def parse(self, request: ParseRequest) -> ParsedDocument:
        path = request.path
        suffix = request.normalized_suffix
        try:
            pages = (
                PyPDFLoader(str(path)).load()
                if suffix == ".pdf"
                else [path.read_text(encoding="utf-8")]
            )
            texts = [
                item.page_content if hasattr(item, "page_content") else item
                for item in pages
            ]
        except Exception as exc:
            raise DocumentParseError("无法解析文档，请确认文件内容和编码正确") from exc
        cleaned = [text.strip() for text in texts if text.strip()]
        if not cleaned:
            raise DocumentParseError()
        combined = "\n\n".join(cleaned)
        warnings = list(("预览已截断",) if len(combined) > 8000 else ())
        page_results: list[dict[str, object]] = []
        if suffix == ".pdf":
            pdf_results, pdf_warnings = self._inspect_pdf_pages(path)
            page_results.extend(pdf_results)
            warnings.extend(pdf_warnings)
        else:
            page_results.append({"page": 1, "kind": "text", "text_chars": len(combined), "image_count": 0})
        counts = {
            kind: sum(item["kind"] == kind for item in page_results)
            for kind in ("text", "table_like", "scanned_or_image")
        }
        deduped_warnings = tuple(dict.fromkeys(warnings))
        elements = tuple(
            ParsedElement(
                element_id=f"element-{index}",
                kind="paragraph",
                text=text,
                page_no=index,
                order=index,
                metadata={"source": "local", "suffix": suffix},
            )
            for index, text in enumerate(cleaned, start=1)
        )
        return ParsedDocument(
            document_metadata={
                "file_name": request.file_name or path.name,
                "suffix": suffix,
                "page_count": len(pages),
                "parser": self.name,
            },
            elements=elements,
            assets=(),
            warnings=deduped_warnings,
            quality=ParseQuality(
                status="warning" if deduped_warnings else "pass",
                counts=counts,
                page_results=tuple(page_results),
                warnings=deduped_warnings,
                parser_name=self.name,
            ),
        )

    @staticmethod
    def _inspect_pdf_pages(
        path: Path,
    ) -> tuple[list[dict[str, object]], list[str]]:
        results: list[dict[str, object]] = []
        warnings: list[str] = []
        try:
            reader = PdfReader(str(path))
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                resources = page.get("/Resources") or {}
                xobjects = (
                    resources.get("/XObject")
                    if hasattr(resources, "get")
                    else None
                )
                image_count = len(xobjects.get_object()) if xobjects else 0
                table_like = "|" in text or any(
                    "  " in line for line in text.splitlines()
                )
                if len(text) < 30 and image_count:
                    kind = "scanned_or_image"
                    warnings.append(
                        f"第{index}页疑似扫描页或图片，文本可能缺失"
                    )
                elif table_like:
                    kind = "table_like"
                    warnings.append(
                        f"第{index}页包含疑似表格，需人工核对结构"
                    )
                else:
                    kind = "text"
                results.append(
                    {
                        "page": index,
                        "kind": kind,
                        "text_chars": len(text),
                        "image_count": image_count,
                    }
                )
        except Exception as exc:
            raise DocumentParseError(
                "PDF逐页质量检查失败，不能安全进入审核"
            ) from exc
        return results, warnings


class DocxStructuredDocumentParser:
    name = "local_docx"

    def parse(self, request: ParseRequest) -> ParsedDocument:
        try:
            document = DocxDocument(str(request.path))
        except Exception as exc:
            raise DocumentParseError("无法解析DOCX文档") from exc
        elements: list[ParsedElement] = []
        order = 0
        for child in document.element.body:
            if child.tag.endswith("}p"):
                paragraph = next(
                    item for item in document.paragraphs if item._p is child
                )
                text = paragraph.text.strip()
                if not text:
                    continue
                order += 1
                style_name = (paragraph.style.name or "").lower()
                is_list = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
                kind = (
                    "title"
                    if style_name.startswith("heading") or style_name.startswith("标题")
                    else "list"
                    if is_list or "list" in style_name
                    else "paragraph"
                )
                elements.append(
                    ParsedElement(
                        element_id=f"docx-{order}",
                        kind=kind,
                        text=text,
                        page_no=None,
                        order=order,
                        metadata={"style": paragraph.style.name or ""},
                    )
                )
            elif child.tag.endswith("}tbl"):
                table = next(item for item in document.tables if item._tbl is child)
                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                    if any(cell.text.strip() for cell in row.cells)
                ]
                if not rows:
                    continue
                order += 1
                text = "\n".join(" | ".join(cell for cell in row) for row in rows)
                html_rows = "".join(
                    "<tr>"
                    + "".join(f"<td>{escape(cell)}</td>" for cell in row)
                    + "</tr>"
                    for row in rows
                )
                elements.append(
                    ParsedElement(
                        element_id=f"docx-{order}",
                        kind="table",
                        text=text,
                        page_no=None,
                        order=order,
                        table_html=f"<table>{html_rows}</table>",
                    )
                )
        if not elements:
            raise DocumentParseError()
        return _build_document(request, self.name, tuple(elements))


class MarkdownStructuredDocumentParser:
    name = "local_markdown"

    def __init__(self) -> None:
        self.markdown = MarkdownIt("commonmark").enable("table")

    def parse(self, request: ParseRequest) -> ParsedDocument:
        text = _read_utf8(request.path)
        tokens = self.markdown.parse(text)
        elements: list[ParsedElement] = []
        order = 0
        index = 0
        while index < len(tokens):
            token = tokens[index]
            if token.type in {"heading_open", "paragraph_open"}:
                inline = tokens[index + 1] if index + 1 < len(tokens) else None
                content = (inline.content if inline and inline.type == "inline" else "").strip()
                if content:
                    order += 1
                    kind = "title" if token.type == "heading_open" else "paragraph"
                    elements.append(
                        ParsedElement(
                            element_id=f"markdown-{order}",
                            kind=kind,
                            text=content,
                            page_no=1,
                            order=order,
                            metadata={"markup": token.tag},
                        )
                    )
            elif token.type == "list_item_open":
                parts: list[str] = []
                depth = token.level
                index += 1
                while index < len(tokens) and not (
                    tokens[index].type == "list_item_close"
                    and tokens[index].level == depth
                ):
                    if tokens[index].type == "inline" and tokens[index].content.strip():
                        parts.append(tokens[index].content.strip())
                    index += 1
                content = " ".join(parts).strip()
                if content:
                    order += 1
                    elements.append(
                        ParsedElement(
                            element_id=f"markdown-{order}",
                            kind="list",
                            text=content,
                            page_no=1,
                            order=order,
                        )
                    )
            elif token.type == "table_open":
                table_text, table_html, index = _consume_markdown_table(tokens, index)
                if table_text:
                    order += 1
                    elements.append(
                        ParsedElement(
                            element_id=f"markdown-{order}",
                            kind="table",
                            text=table_text,
                            page_no=1,
                            order=order,
                            table_html=table_html,
                        )
                    )
            index += 1
        if not elements:
            raise DocumentParseError()
        return _build_document(request, self.name, tuple(elements), page_count=1)


class HtmlStructuredDocumentParser:
    name = "local_html"
    dangerous_tags = {"script", "style", "form", "iframe", "object", "embed", "noscript"}

    def parse(self, request: ParseRequest) -> ParsedDocument:
        soup = BeautifulSoup(_read_utf8(request.path), "html.parser")
        for tag in soup.find_all(self.dangerous_tags):
            tag.decompose()
        for tag in soup.find_all(True):
            tag.attrs = {}
        root = soup.body or soup
        elements: list[ParsedElement] = []
        order = 0
        for node in root.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
            if any(parent.name == "table" for parent in node.parents) and node.name != "table":
                continue
            text = node.get_text(" ", strip=True)
            if not text:
                continue
            order += 1
            kind = (
                "title"
                if node.name and node.name.startswith("h")
                else "list"
                if node.name == "li"
                else "table"
                if node.name == "table"
                else "paragraph"
            )
            elements.append(
                ParsedElement(
                    element_id=f"html-{order}",
                    kind=kind,
                    text=text,
                    page_no=1,
                    order=order,
                    table_html=str(node) if kind == "table" else None,
                    metadata={"tag": node.name or ""},
                )
            )
        if not elements:
            raise DocumentParseError()
        return _build_document(request, self.name, tuple(elements), page_count=1)


class ImageStructuredDocumentParser:
    name = "local_image"

    def parse(self, request: ParseRequest) -> ParsedDocument:
        try:
            with Image.open(request.path) as image:
                width, height = image.size
                mime_type = image.get_format_mimetype() or FileTypePolicy.mime_type_for_suffix(
                    request.normalized_suffix
                )
        except (UnidentifiedImageError, OSError) as exc:
            raise DocumentParseError("Invalid image document") from exc
        data = request.path.read_bytes()
        digest = hashlib.sha256(data).hexdigest()
        byte_size = len(data)
        pixel_count = width * height
        file_name = request.file_name or request.path.name
        warning = "Image document is waiting for OCR/Vision enrichment before publication."
        asset = ParsedAsset(
            asset_id="image-1",
            kind="uploaded_image",
            page_no=1,
            mime_type=mime_type,
            storage_ref="discovered://uploaded-image",
            sha256=digest,
            metadata={
                "file_name": file_name,
                "width": width,
                "height": height,
                "pixel_count": pixel_count,
                "byte_size": byte_size,
                "purpose": "document_understanding",
                "source_kind": "uploaded_image_file",
                "materialized": False,
            },
        )
        return ParsedDocument(
            document_metadata={
                "file_name": file_name,
                "suffix": request.normalized_suffix,
                "page_count": 1,
                "parser": self.name,
                "enrichment": {
                    "status": "waiting_enrichment",
                    "reason": "image_input",
                    "asset_count": 1,
                },
            },
            elements=(),
            assets=(asset,),
            warnings=(warning,),
            quality=ParseQuality(
                status="warning",
                counts={
                    "text": 0,
                    "table_like": 0,
                    "scanned_or_image": 1,
                    "asset": 1,
                },
                page_results=(
                    {
                        "page": 1,
                        "kind": "scanned_or_image",
                        "text_chars": 0,
                        "image_count": 1,
                        "width": width,
                        "height": height,
                    },
                ),
                warnings=(warning,),
                parser_name=self.name,
            ),
        )


def _read_utf8(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("文本文件必须使用UTF-8编码") from exc
    if not text.strip():
        raise DocumentParseError()
    return text


def _build_document(
    request: ParseRequest,
    parser_name: str,
    elements: tuple[ParsedElement, ...],
    *,
    page_count: int | None = None,
    warnings: tuple[str, ...] = (),
) -> ParsedDocument:
    counts = {
        kind: sum(element.kind == kind for element in elements)
        for kind in ("title", "paragraph", "list", "table")
    }
    return ParsedDocument(
        document_metadata={
            "file_name": request.file_name or request.path.name,
            "suffix": request.normalized_suffix,
            "page_count": page_count or max(
                (element.page_no or 0 for element in elements), default=1
            ),
            "parser": parser_name,
        },
        elements=elements,
        assets=(),
        warnings=warnings,
        quality=ParseQuality(
            status="warning" if warnings else "pass",
            counts=counts,
            page_results=(),
            warnings=warnings,
            parser_name=parser_name,
        ),
    )


def _consume_markdown_table(tokens, start: int) -> tuple[str, str, int]:
    rows: list[list[str]] = []
    current_row: list[str] | None = None
    index = start + 1
    while index < len(tokens) and tokens[index].type != "table_close":
        token = tokens[index]
        if token.type == "tr_open":
            current_row = []
        elif token.type == "tr_close":
            if current_row and any(current_row):
                rows.append(current_row)
            current_row = None
        elif token.type == "inline" and current_row is not None:
            current_row.append(token.content.strip())
        index += 1
    text = "\n".join(" | ".join(cell for cell in row) for row in rows)
    html_rows = "".join(
        "<tr>" + "".join(f"<td>{escape(cell)}</td>" for cell in row) + "</tr>"
        for row in rows
    )
    return text, f"<table>{html_rows}</table>" if rows else "", index


def _with_warning(document: ParsedDocument, warning: str) -> ParsedDocument:
    warnings = tuple(dict.fromkeys((*document.warnings, warning)))
    return ParsedDocument(
        document_metadata=document.document_metadata,
        elements=document.elements,
        assets=document.assets,
        warnings=warnings,
        quality=ParseQuality(
            status="warning",
            counts=document.quality.counts,
            page_results=document.quality.page_results,
            warnings=warnings,
            parser_name=document.quality.parser_name,
        ),
    )


def build_default_parser_registry(
    *,
    pdf_candidate: DocumentParserPort | None = None,
    pdf_candidate_enabled: bool = False,
    pdf_candidate_promoted: bool = False,
    pdf_candidate_max_pages: int = 20,
    pdf_candidate_max_file_size_bytes: int = 10 * 1024 * 1024,
) -> ParserRegistry:
    local_pdf_txt = LocalStructuredDocumentParser()
    markdown = MarkdownStructuredDocumentParser()
    pdf = PdfCandidateFallbackParser(
        local_pdf_txt,
        pdf_candidate,
        enabled=pdf_candidate_enabled,
        promoted=pdf_candidate_promoted,
        max_pages=pdf_candidate_max_pages,
        max_file_size_bytes=pdf_candidate_max_file_size_bytes,
    )
    return ParserRegistry(
        [
            ParserRegistration(
                name=LocalStructuredDocumentParser.name,
                parser=pdf,
                suffixes=(".pdf", ".txt"),
                mime_types=("application/pdf", "text/plain"),
            ),
            ParserRegistration(
                name=DocxStructuredDocumentParser.name,
                parser=DocxStructuredDocumentParser(),
                suffixes=(".docx",),
                mime_types=(FileTypePolicy.mime_type_for_suffix(".docx"),),
            ),
            ParserRegistration(
                name=MarkdownStructuredDocumentParser.name,
                parser=markdown,
                suffixes=(".md", ".markdown"),
                mime_types=("text/markdown",),
            ),
            ParserRegistration(
                name=HtmlStructuredDocumentParser.name,
                parser=HtmlStructuredDocumentParser(),
                suffixes=(".html", ".htm"),
                mime_types=("text/html",),
            ),
            ParserRegistration(
                name=ImageStructuredDocumentParser.name,
                parser=ImageStructuredDocumentParser(),
                suffixes=(".png", ".jpg", ".jpeg"),
                mime_types=("image/png", "image/jpeg"),
            ),
        ]
    )


class LocalDocumentParser:
    def __init__(
        self,
        registry: ParserRegistry | None = None,
        *,
        pdf_candidate: DocumentParserPort | None = None,
        pdf_candidate_enabled: bool = False,
        pdf_candidate_promoted: bool = False,
        pdf_candidate_max_pages: int = 20,
        pdf_candidate_max_file_size_bytes: int = 10 * 1024 * 1024,
    ) -> None:
        self.registry = registry or build_default_parser_registry(
            pdf_candidate=pdf_candidate,
            pdf_candidate_enabled=pdf_candidate_enabled,
            pdf_candidate_promoted=pdf_candidate_promoted,
            pdf_candidate_max_pages=pdf_candidate_max_pages,
            pdf_candidate_max_file_size_bytes=pdf_candidate_max_file_size_bytes,
        )

    def parse_document(self, request: ParseRequest) -> ParsedDocument:
        return self.registry.parse(request)

    def parse(self, path: Path, suffix: str) -> ParsedPreview:
        document = self.parse_document(
            ParseRequest(path=path, suffix=suffix, file_name=path.name)
        )
        return ParsedPreview.from_document(document)
